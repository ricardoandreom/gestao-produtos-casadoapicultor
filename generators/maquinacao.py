"""Gerador do documento de MAQUINAÇÃO (Ordem de Produção).

Estratégia: abre o template 'maq_temp.docx' como base, o que preserva
exactamente o cabeçalho de página (logo VML + tabela de datas flutuante)
e as definições da página (margens, estilos).

Em seguida:
  1. Substitui o número de ordem na célula (0,2) do header
  2. Remove apenas os parágrafos do body (a tabela de datas fica intacta)
  3. Reconstrói os parágrafos com os dados da encomenda

Valores medidos directamente do template (XML):
  - Título:   Arial Black 38 pt, bold, underline, centrado, ind=-993 twips
  - Categoria: Arial Black 26 pt, bold, underline, ind=-993 twips, sa=0
  - Produto:   Arial Black 26 pt, ind=-993 twips, sa=6 pt
  - Notas:     Arial Black 26 pt, ind=-993 twips, sa=6 pt  (prefixo "- ")
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ._common import set_font, carregar_mapping_completo


TEMPLATE_PATH = Path(__file__).parent.parent / "data" / "doc_templates" / "maq_temp.docx"
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Indentação negativa usada em todos os parágrafos do body (twips)
_INDENT = -993


# ---------------------------------------------------------------------------
# Helpers internos
# ---------------------------------------------------------------------------

def _set_left_indent_twips(p, twips: int):
    """Define w:ind/@w:left directamente em twips (evita conversão EMU)."""
    pPr = p._p.get_or_add_pPr()
    # Remove qualquer <w:ind> existente para evitar duplicados
    for old in pPr.findall(f"{{{_W}}}ind"):
        pPr.remove(old)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(twips))
    pPr.append(ind)


def _body_para(doc, text="", bold=False, underline=False, size_pt=26,
               align=None, space_after_pt=0, space_before_pt=0):
    """Adiciona um parágrafo ao body com as propriedades do template."""
    p = doc.add_paragraph()

    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after  = Pt(space_after_pt)
    _set_left_indent_twips(p, _INDENT)

    if text:
        run = p.add_run(text)
        if underline:
            run.underline = True
        set_font(run, font_name="Arial Black", size_pt=size_pt, bold=bold)

    return p


# ---------------------------------------------------------------------------
# Actualização do número de ordem no header
# ---------------------------------------------------------------------------

def _update_header_order_num(doc, encomenda):
    """Substitui o texto do número de ordem na célula (0,2) do header."""
    _day, _month, year = encomenda["data"].split("/")
    order_num = f" {encomenda['id']:03d}/{year[2:]} {encomenda['cliente_tipo']}"

    header = doc.sections[0].header
    cell   = header.tables[0].cell(0, 2)
    para   = cell.paragraphs[0]

    # Apaga todos os runs existentes
    for run in list(para.runs):
        run._r.getparent().remove(run._r)

    # Insere o novo número numa só run
    run = para.add_run(order_num)
    run.bold = True
    set_font(run, font_name="Arial", size_pt=20, bold=True)


# ---------------------------------------------------------------------------
# Limpeza e reconstrução do body
# ---------------------------------------------------------------------------

def _make_date_table_inline(doc):
    """Converte a tabela flutuante de datas em tabela inline alinhada com o cabeçalho.

    Header: tblW=10632, tblInd=-1000 → borda direita a 9632 twips do texto.
    Dates table: tblW=4819 → tblInd = 9632 - 4819 = 4813 para alinhar bordas direitas.
    """
    body = doc.element.body
    for tbl in body.findall(f"{{{_W}}}tbl"):
        tblPr = tbl.find(f"{{{_W}}}tblPr")
        if tblPr is None:
            continue
        # Remove floating positioning
        for tag in (f"{{{_W}}}tblpPr", f"{{{_W}}}tblOverlap"):
            el = tblPr.find(tag)
            if el is not None:
                tblPr.remove(el)
        # Keep original width (4819 dxa)
        tblW = tblPr.find(f"{{{_W}}}tblW")
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.insert(0, tblW)
        tblW.set(qn("w:w"), "4819")
        tblW.set(qn("w:type"), "dxa")
        # Shift left edge so right edge aligns with header's right border
        tblInd = tblPr.find(f"{{{_W}}}tblInd")
        if tblInd is None:
            tblInd = OxmlElement("w:tblInd")
            tblPr.append(tblInd)
        tblInd.set(qn("w:w"), "4813")
        tblInd.set(qn("w:type"), "dxa")
        # Left-aligned so tblInd positions correctly
        jc = tblPr.find(f"{{{_W}}}jc")
        if jc is None:
            jc = OxmlElement("w:jc")
            tblPr.append(jc)
        jc.set(qn("w:val"), "left")


def _clear_body_paragraphs(doc):
    """Remove todos os parágrafos do body, preservando tabelas e sectPr."""
    body = doc.element.body
    for child in [c for c in body if c.tag == f"{{{_W}}}p"]:
        body.remove(child)


def _agrupar_por_categoria(produtos, mapping):
    """Agrupa produtos por categoria, preservando a ordem de entrada."""
    grupos: dict[str, list] = {}
    for produto in produtos:
        nome_up   = produto["nome"].strip().upper()
        info      = mapping.get(nome_up, {})
        categoria = info.get("categoria", "").strip() or "OUTROS"
        notas_map = info.get("notas", "").strip()
        grupos.setdefault(categoria, []).append((produto, notas_map))
    return grupos


def _add_body(doc, encomenda, mapping):
    """Reconstrói os parágrafos do body com os dados da encomenda."""
    grupos = _agrupar_por_categoria(encomenda["produtos"], mapping)

    # Parágrafo vazio antes do título (espaçamento acima do título)
    _body_para(doc, "")

    # Título "FAZER – PRODUÇÃO {tipo}"
    _body_para(
        doc,
        f"FAZER – PRODUÇÃO {encomenda['cliente_tipo']}",
        bold=True, underline=True, size_pt=38,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    # Categorias e produtos
    first = True
    for categoria, itens in grupos.items():
        # Linha em branco entre categorias
        if not first:
            _body_para(doc, "", space_after_pt=6)
        first = False

        # Cabeçalho da categoria (sublinhado)
        _body_para(doc, categoria, bold=True, underline=True, size_pt=26)

        for produto, notas_map in itens:
            # Linha do produto: "{qty} – {nome}"
            _body_para(
                doc,
                f"{produto['quantidade']} – {produto['nome']}",
                size_pt=26, space_after_pt=6,
            )

            # Notas: as do utilizador sobrepõem as do mapping
            notas = produto.get("notas", "").strip() or notas_map
            if notas:
                _body_para(doc, f"- {notas}", size_pt=26, space_after_pt=6)


# ---------------------------------------------------------------------------
# Ponto de entrada
# ---------------------------------------------------------------------------

def gerar(encomenda: dict, caminho_saida: str):
    """Gera o documento de maquinação para a encomenda."""
    mapping = carregar_mapping_completo()

    doc = Document(str(TEMPLATE_PATH))

    _update_header_order_num(doc, encomenda)
    _make_date_table_inline(doc)
    _clear_body_paragraphs(doc)
    _add_body(doc, encomenda, mapping)

    doc.save(caminho_saida)
