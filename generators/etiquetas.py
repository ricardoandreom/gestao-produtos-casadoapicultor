"""Gerador do documento de ETIQUETAS.

Abre o template 'etiquetas_temp.docx' e, para cada produto da encomenda,
escreve N vezes a linha:
    {qt_palete} - {nome_produto}

N = ceil(quantidade / qt_palete).
Se o produto não tiver mapping, escreve a linha uma única vez sem a quantidade.
"""

import math
from pathlib import Path

from docx import Document
from docx.shared import Pt

from ._common import set_font, carregar_mapping, qt_palete


TEMPLATE_PATH = Path(__file__).parent.parent / "data" / "doc_templates" / "etiquetas_temp.docx"
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _clear_body_paragraphs(doc):
    body = doc.element.body
    for child in [c for c in body if c.tag == f"{{{_W}}}p"]:
        body.remove(child)


def _add_linha(doc, texto: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(texto)
    set_font(run, font_name="Arial Black", size_pt=48, bold=True)


def gerar(encomenda: dict, caminho_saida: str):
    """Gera o documento de etiquetas para a encomenda."""
    mapping = carregar_mapping()
    cliente_tipo = encomenda.get("cliente_tipo", "AM")

    doc = Document(str(TEMPLATE_PATH))
    _clear_body_paragraphs(doc)

    for produto in encomenda["produtos"]:
        nome_upper = produto["nome"].strip().upper()
        qt_str = mapping.get(nome_upper, "")
        qt_str_cliente = qt_palete(qt_str, cliente_tipo)

        if qt_str_cliente:
            try:
                qt_val = float(qt_str_cliente)
            except ValueError:
                qt_val = None
        else:
            qt_val = None

        if qt_val:
            try:
                quantidade = float(str(produto.get("quantidade", 1)).replace(",", "."))
            except ValueError:
                quantidade = 1.0
            repeticoes = math.ceil(quantidade / qt_val)
            linha = f"{qt_str_cliente} - {produto['nome']}"
        else:
            repeticoes = 1
            linha = produto["nome"]

        for _ in range(repeticoes):
            _add_linha(doc, linha)

        # Linha em branco entre produtos
        _add_linha(doc, "")

    doc.save(caminho_saida)
