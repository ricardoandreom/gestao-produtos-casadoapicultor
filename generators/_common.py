"""
Helpers partilhados pelos geradores de documentos Word.

Aqui ficam todas as funções utilitárias que mais do que um gerador usa:
- formatação de fontes/parágrafos
- linha horizontal
- carregamento do mapping de produtos (products_mapping.xlsx)
- extração da quantidade por palete consoante o tipo de cliente (AM/FR/OUTRO)
- helpers de construção de documentos Word (parágrafos, tabela de datas, cabeçalho)
"""

from pathlib import Path

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --- Constantes Word partilhadas ---

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_INDENT = -993  # twips — expansão para a margem esquerda (igual em todos os geradores)


# --- Formatação de texto ---

def set_font(run, font_name="Arial Black", size_pt=22, bold=False, color=None):
    """Aplica fonte/tamanho/cor a um run e força a fonte também ao nível do XML
    (necessário em alguns Words que ignoram a propriedade Python)."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.insert(0, rFonts)


def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT,
                  font_name="Arial Black", size_pt=22, bold=False,
                  space_before=0, space_after=200, color=None):
    """Adiciona um parágrafo formatado ao documento e devolve-o."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, font_name=font_name, size_pt=size_pt, bold=bold, color=color)
    return p


def add_horizontal_line(doc):
    """Insere uma linha horizontal cinzenta no documento."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


# --- Mapping de produtos ---

def carregar_mapping():
    """Carrega o products_mapping.xlsx e devolve um dict
    {nome_produto_em_uppercase: qt_palete_AM_FR}."""
    import pandas as pd

    # data/ está dois níveis acima deste ficheiro: generators/ -> raiz -> data/
    mapping_path = Path(__file__).parent.parent / "data" / "products_mapping.xlsx"
    if not mapping_path.exists():
        return {}

    df = pd.read_excel(str(mapping_path), dtype=str).fillna("")
    df.columns = [c.strip() for c in df.columns]

    mapping = {}
    for _, row in df.iterrows():
        nome = str(row.get("nome_produto", "")).strip().upper()
        qt   = str(row.get("qt_palete_AM_FR", "")).strip()
        if nome:
            mapping[nome] = qt
    return mapping


def carregar_mapping_completo():
    """Carrega o products_mapping.xlsx e devolve um dict
    {nome_produto_em_uppercase: {"categoria": str, "notas": str}}."""
    import pandas as pd

    mapping_path = Path(__file__).parent.parent / "data" / "products_mapping.xlsx"
    if not mapping_path.exists():
        return {}

    df = pd.read_excel(str(mapping_path), dtype=str).fillna("")
    df.columns = [c.strip() for c in df.columns]

    mapping = {}
    for _, row in df.iterrows():
        nome = str(row.get("nome_produto", "")).strip().upper()
        if nome:
            mapping[nome] = {
                "categoria": str(row.get("categoria_produto", "")).strip().upper(),
                "notas":     str(row.get("notas", "")).strip(),
            }
    return mapping


def carregar_mapping_acabamento():
    """Carrega o products_mapping.xlsx e devolve um dict
    {nome_produto_em_uppercase: {ref_interna, qt_palete_AM_FR, notas, nome_produto}}."""
    import pandas as pd

    mapping_path = Path(__file__).parent.parent / "data" / "products_mapping.xlsx"
    if not mapping_path.exists():
        return {}

    df = pd.read_excel(str(mapping_path), dtype=str).fillna("")
    df.columns = [c.strip() for c in df.columns]

    mapping = {}
    for _, row in df.iterrows():
        nome = str(row.get("nome_produto", "")).strip().upper()
        if nome:
            mapping[nome] = {
                "nome_produto":   str(row.get("nome_produto", "")).strip(),
                "ref_interna":    str(row.get("ref_interna", "")).strip(),
                "qt_palete_AM_FR": str(row.get("qt_palete_AM_FR", "")).strip(),
                "notas":          str(row.get("notas", "")).strip(),
            }
    return mapping


# --- Helpers de construção de documentos Word ---

def _set_left_indent_twips(p, twips: int):
    """Define w:ind/@w:left directamente em twips."""
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(f"{{{_W}}}ind"):
        pPr.remove(old)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(twips))
    pPr.append(ind)


def _body_para(doc, text="", bold=False, underline=False, size_pt=26,
               align=None, space_after_pt=0, space_before_pt=0):
    """Adiciona um parágrafo ao body com indentação e fonte padrão dos geradores."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after  = Pt(space_after_pt)
    _set_left_indent_twips(p, _INDENT)
    if text:
        run = p.add_run(text)
        if underline:
            run.underline = True
        set_font(run, font_name="Arial Black", size_pt=size_pt, bold=bold)
    return p


def _update_header_order_num(doc, encomenda):
    """Substitui o número de ordem na célula (0,2) do header."""
    _day, _month, year = encomenda["data"].split("/")
    order_num = f" {encomenda['id']:03d}/{year[2:]} {encomenda['cliente_tipo']}"

    header = doc.sections[0].header
    cell   = header.tables[0].cell(0, 2)
    para   = cell.paragraphs[0]

    for run in list(para.runs):
        run._r.getparent().remove(run._r)

    run = para.add_run(order_num)
    run.bold = True
    set_font(run, font_name="Arial", size_pt=20, bold=True)


def _make_date_table_inline(doc):
    """Converte a tabela flutuante de datas em inline, alinhando a borda direita
    com a borda direita do cabeçalho (tblInd=4813 para tblW=4819, header=10632/-1000).
    """
    body = doc.element.body
    for tbl in body.findall(f"{{{_W}}}tbl"):
        tblPr = tbl.find(f"{{{_W}}}tblPr")
        if tblPr is None:
            continue
        for tag in (f"{{{_W}}}tblpPr", f"{{{_W}}}tblOverlap"):
            el = tblPr.find(tag)
            if el is not None:
                tblPr.remove(el)
        tblW = tblPr.find(f"{{{_W}}}tblW")
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.insert(0, tblW)
        tblW.set(qn("w:w"), "4819")
        tblW.set(qn("w:type"), "dxa")
        tblInd = tblPr.find(f"{{{_W}}}tblInd")
        if tblInd is None:
            tblInd = OxmlElement("w:tblInd")
            tblPr.append(tblInd)
        tblInd.set(qn("w:w"), "4813")
        tblInd.set(qn("w:type"), "dxa")
        jc = tblPr.find(f"{{{_W}}}jc")
        if jc is None:
            jc = OxmlElement("w:jc")
            tblPr.append(jc)
        jc.set(qn("w:val"), "left")


def _clear_body_paragraphs(doc):
    """Remove todos os parágrafos do body, preservando tabelas e sectPr."""
    body = doc.element.body
    for child in [c for c in body if c.tag == f"{{{_W}}}p"]:
        body.remove(child)


# --- Mapping de produtos ---

def qt_palete(qt_str: str, cliente_tipo: str) -> str:
    """Da string '20/30' (AM/FR), devolve a parte que corresponde ao cliente_tipo.

    - AM e OUTRO → parte ANTES da "/"
    - FR         → parte DEPOIS da "/"
    """
    if not qt_str:
        return ""
    partes = qt_str.split("/")
    if cliente_tipo == "FR":
        return partes[1].strip() if len(partes) > 1 else partes[0].strip()
    # AM, OUTRO, ou qualquer outro valor desconhecido → parte AM
    return partes[0].strip()