"""Gerador do documento de ACABAMENTO (Montagem e Acabamento).

Estratégia idêntica à da maquinação: abre acabamento_temp.docx como base,
que já tem o cabeçalho com o logo e as 3 linhas de datas no body.

Conteúdo por encomenda:
  Título:  FAZER - {total paletes} PALETES - {tipo}

  Por produto:
    {n paletes} - PALETES - CADA COM {qt_palete}
    - REF: {ref_interna}
    - {nome_produto}
    - {notas}          (se existirem)
    - TOTAL {quantidade}
"""

import math
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ._common import set_font, carregar_mapping_acabamento, qt_palete


TEMPLATE_PATH = Path(__file__).parent.parent / "data" / "doc_templates" / "acabamento_temp.docx"
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_INDENT = -993  # twips — idêntico ao maquinação, expande texto para a margem esquerda


# ---------------------------------------------------------------------------
# Helpers internos
# ---------------------------------------------------------------------------

def _set_left_indent_twips(p, twips: int):
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(f"{{{_W}}}ind"):
        pPr.remove(old)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(twips))
    pPr.append(ind)


def _body_para(doc, text="", bold=False, underline=False, size_pt=26,
               align=None, space_after_pt=0, space_before_pt=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after  = Pt(space_after_pt)
    _set_left_indent_twips(p, _INDENT)
    if text:
        run = p.add_run(text)
        if underline:
            run.underline = True
        set_font(run, font_name="Arial Black", size_pt=size_pt, bold=bold)
    return p


# ---------------------------------------------------------------------------
# Actualização do cabeçalho
# ---------------------------------------------------------------------------

def _update_header_order_num(doc, encomenda):
    """Substitui o número de ordem na célula (0,2) do header."""
    _day, _month, year = encomenda["data"].split("/")
    order_num = f" {encomenda['id']:03d}/{year[2:]} {encomenda['cliente_tipo']}"

    header = doc.sections[0].header
    cell   = header.tables[0].cell(0, 2)
    para   = cell.paragraphs[0]

    for run in list(para.runs):
        run._r.getparent().remove(run._r)

    run = para.add_run(order_num)
    run.bold = True
    set_font(run, font_name="Arial", size_pt=20, bold=True)


# ---------------------------------------------------------------------------
# Tabela de datas: converter de flutuante para inline à direita
# ---------------------------------------------------------------------------

def _make_date_table_inline(doc):
    # Header table: tblW=10632, tblInd=-1000 → right edge at 9632 twips from text left.
    # Dates table: tblW=4819 → tblInd = 9632 - 4819 = 4813 to align right edges.
    body = doc.element.body
    for tbl in body.findall(f"{{{_W}}}tbl"):
        tblPr = tbl.find(f"{{{_W}}}tblPr")
        if tblPr is None:
            continue
        for tag in (f"{{{_W}}}tblpPr", f"{{{_W}}}tblOverlap"):
            el = tblPr.find(tag)
            if el is not None:
                tblPr.remove(el)
        # Keep original width (4819 dxa)
        tblW = tblPr.find(f"{{{_W}}}tblW")
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.insert(0, tblW)
        tblW.set(qn("w:w"), "4819")
        tblW.set(qn("w:type"), "dxa")
        # Shift left edge so right edge aligns with header's right border
        tblInd = tblPr.find(f"{{{_W}}}tblInd")
        if tblInd is None:
            tblInd = OxmlElement("w:tblInd")
            tblPr.append(tblInd)
        tblInd.set(qn("w:w"), "4813")
        tblInd.set(qn("w:type"), "dxa")
        # Ensure left-aligned (so tblInd positions correctly)
        jc = tblPr.find(f"{{{_W}}}jc")
        if jc is None:
            jc = OxmlElement("w:jc")
            tblPr.append(jc)
        jc.set(qn("w:val"), "left")


# ---------------------------------------------------------------------------
# Limpar parágrafos do body (mantém tabelas e sectPr)
# ---------------------------------------------------------------------------

def _clear_body_paragraphs(doc):
    body = doc.element.body
    for child in [c for c in body if c.tag == f"{{{_W}}}p"]:
        body.remove(child)


# ---------------------------------------------------------------------------
# Cálculo de paletes
# ---------------------------------------------------------------------------

def _calcular_n_paletes(quantidade: int, qt_str: str, cliente_tipo: str):
    """Devolve (n_paletes int, qt_val_str str) ou (None, qt_str) se não calculável."""
    qt_val_str = qt_palete(qt_str, cliente_tipo)
    try:
        qt_val = int(qt_val_str)
        if qt_val <= 0:
            return None, qt_val_str
        return math.ceil(quantidade / qt_val), qt_val_str
    except (ValueError, TypeError):
        return None, qt_val_str or qt_str


# ---------------------------------------------------------------------------
# Construção do body
# ---------------------------------------------------------------------------

def _add_body(doc, encomenda, mapping):
    tipo     = encomenda["cliente_tipo"]
    produtos = encomenda["produtos"]

    # Preparar dados por produto
    items = []
    for produto in produtos:
        nome_up = produto["nome"].strip().upper()
        info    = mapping.get(nome_up, {})
        qt_str  = info.get("qt_palete_AM_FR", "")
        n_paletes, qt_val_str = _calcular_n_paletes(
            int(produto["quantidade"]), qt_str, tipo
        )
        items.append({
            "produto":    produto,
            "info":       info,
            "n_paletes":  n_paletes,
            "qt_val_str": qt_val_str,
        })

    total_paletes = sum(i["n_paletes"] for i in items if i["n_paletes"] is not None)

    # Parágrafo vazio de espaçamento antes do título
    _body_para(doc, "")

    # Título
    _body_para(
        doc,
        f"FAZER - {total_paletes} PALETES - {tipo}",
        bold=True, underline=True, size_pt=38,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    # Produtos
    for i, item in enumerate(items):
        if i > 0:
            _body_para(doc, "", space_after_pt=6)

        produto    = item["produto"]
        info       = item["info"]
        n_paletes  = item["n_paletes"]
        qt_display = item["qt_val_str"] or "?"
        n_str      = str(n_paletes) if n_paletes is not None else "?"

        # Linha 1: "{n} - PALETES - CADA COM {qt}"
        _body_para(
            doc,
            f"{n_str} - PALETES - CADA COM {qt_display}",
            bold=True, size_pt=26, space_after_pt=0,
        )

        # Linha 2: "- REF: {ref}"
        ref = info.get("ref_interna", "").strip()
        _body_para(doc, f"- REF: {ref}" if ref else "- REF: —", size_pt=26, space_after_pt=0)

        # Linha 3: "- {nome_produto do mapping ou o que foi digitado}"
        nome_map = info.get("nome_produto", "").strip() or produto["nome"]
        _body_para(doc, f"- {nome_map}", size_pt=26, space_after_pt=0)

        # Linha 4: "- {notas}" (notas do utilizador primeiro; se vazias, as do mapping)
        notas = produto.get("notas", "").strip() or info.get("notas", "").strip()
        if notas:
            _body_para(doc, f"- {notas}", size_pt=26, space_after_pt=0)

        # Linha 5: "- TOTAL {quantidade}"
        _body_para(doc, f"- TOTAL {produto['quantidade']}", size_pt=26, space_after_pt=6)


# ---------------------------------------------------------------------------
# Ponto de entrada
# ---------------------------------------------------------------------------

def gerar(encomenda: dict, caminho_saida: str):
    """Gera o documento de acabamento para a encomenda."""
    mapping = carregar_mapping_acabamento()

    doc = Document(str(TEMPLATE_PATH))

    _update_header_order_num(doc, encomenda)
    _make_date_table_inline(doc)
    _clear_body_paragraphs(doc)
    _add_body(doc, encomenda, mapping)

    doc.save(caminho_saida)
