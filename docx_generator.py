from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, font_name="Arial Black", size_pt=22, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Forçar fonte também no XML (compatibilidade)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.insert(0, rFonts)


def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, font_name="Arial Black", size_pt=22, bold=False, space_before=0, space_after=200, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, font_name=font_name, size_pt=size_pt, bold=bold, color=color)
    return p


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _carregar_mapping():
    """Carrega o products_mapping.xlsx e devolve um dict nome_produto -> qt_palete_AM_FR."""
    from pathlib import Path
    import pandas as pd

    mapping_path = Path(__file__).parent / "data" / "products_mapping.xlsx"
    if not mapping_path.exists():
        return {}

    df = pd.read_excel(str(mapping_path), dtype=str).fillna("")
    # Normalizar nomes das colunas (remove espaços extra)
    df.columns = [c.strip() for c in df.columns]
    # Construir dict: nome_produto (uppercase stripped) -> qt_palete_AM_FR
    mapping = {}
    for _, row in df.iterrows():
        nome = str(row.get("nome_produto", "")).strip().upper()
        qt   = str(row.get("qt_palete_AM_FR", "")).strip()
        if nome:
            mapping[nome] = qt
    return mapping


def _qt_palete(qt_str: str, cliente_tipo: str) -> str:
    """Extrai a quantidade AM (antes do /) ou FR (depois do /) da string."""
    if not qt_str:
        return ""
    partes = qt_str.split("/")
    if cliente_tipo == "AM":
        return partes[0].strip()
    else:
        return partes[1].strip() if len(partes) > 1 else partes[0].strip()


def gerar_etiquetas(encomenda: dict, caminho_saida: str):
    """
    Para cada produto da encomenda escreve:
        {quantidade_palete} - {nome_produto}
    A quantidade_palete vem do mapping AM/FR.
    Sem cabeçalhos, sem decoração — só o conteúdo.
    """
    mapping = _carregar_mapping()
    cliente_tipo = encomenda.get("cliente_tipo", "AM")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    for produto in encomenda["produtos"]:
        nome_upper = produto["nome"].strip().upper()
        qt_str = mapping.get(nome_upper, "")
        qt_palete = _qt_palete(qt_str, cliente_tipo)

        if qt_palete:
            linha = f"{qt_palete} - {produto['nome']}"
        else:
            # Se não encontrar no mapping, omite a quantidade de palete
            linha = produto["nome"]

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(linha)
        set_font(run, font_name="Arial Black", size_pt=22, bold=True)

    doc.save(caminho_saida)


def gerar_acabamento(encomenda: dict, caminho_saida: str):
    """Por agora é uma cópia da maquinação."""
    gerar_documento_word(encomenda, caminho_saida)


def gerar_documento_word(encomenda: dict, caminho_saida: str):
    doc = Document()

    # Margens
    section = doc.sections[0]
    section.top_margin = Pt(60)
    section.bottom_margin = Pt(60)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    # Cabeçalho - Título
    add_paragraph(
        doc,
        "ORDEM DE FABRICO",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        font_name="Arial Black",
        size_pt=26,
        bold=True,
        space_before=0,
        space_after=6,
        color=(30, 30, 30)
    )

    add_horizontal_line(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Número de encomenda
    add_paragraph(
        doc,
        f"Encomenda Nº {encomenda['id']:04d}",
        font_name="Arial Black",
        size_pt=22,
        space_after=6
    )

    # Data
    add_paragraph(
        doc,
        f"Data: {encomenda['data']}",
        font_name="Arial Black",
        size_pt=22,
        space_after=6
    )

    # Cliente e tipo
    add_paragraph(
        doc,
        f"Cliente: {encomenda['cliente']}",
        font_name="Arial Black",
        size_pt=22,
        space_after=6
    )

    add_paragraph(
        doc,
        f"Tipo: {encomenda['cliente_tipo']}",
        font_name="Arial Black",
        size_pt=22,
        space_after=16
    )

    add_horizontal_line(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Título produtos
    add_paragraph(
        doc,
        "PRODUTOS A MANUFATURAR",
        font_name="Arial Black",
        size_pt=22,
        bold=True,
        space_after=10
    )

    # Produtos
    for produto in encomenda["produtos"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(20)

        run_bullet = p.add_run("▸  ")
        set_font(run_bullet, font_name="Arial Black", size_pt=22)

        run_prod = p.add_run(f"{produto['nome']}")
        set_font(run_prod, font_name="Arial Black", size_pt=22)

        run_qty = p.add_run(f"   ×{produto['quantidade']}")
        set_font(run_qty, font_name="Arial Black", size_pt=22, bold=True)

        # Notas do produto (opcional)
        notas = produto.get("notas", "").strip()
        if notas:
            p_notas = doc.add_paragraph()
            p_notas.paragraph_format.space_before = Pt(0)
            p_notas.paragraph_format.space_after = Pt(8)
            p_notas.paragraph_format.left_indent = Pt(44)
            run_notas = p_notas.add_run(f"↳ {notas}")
            set_font(run_notas, font_name="Arial Black", size_pt=14, color=(120, 116, 100))
        else:
            p.paragraph_format.space_after = Pt(8)

    add_horizontal_line(doc)

    # Total
    total = sum(p["quantidade"] for p in encomenda["produtos"])
    add_paragraph(
        doc,
        f"Total de unidades: {total}",
        font_name="Arial Black",
        size_pt=22,
        bold=True,
        space_before=10,
        space_after=0
    )

    doc.save(caminho_saida)